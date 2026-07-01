from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "Templates" / "InvoiceTemplateExample.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = table._tbl.tblGrid
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, RGBColor(0x2E, 0x74, 0xB5), 16, 8),
        ("Heading 2", 13, RGBColor(0x2E, 0x74, 0xB5), 12, 6),
        ("Heading 3", 12, RGBColor(0x1F, 0x4D, 0x78), 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_text(paragraph, text, bold=False, size=None, color=None):
    run = paragraph.add_run(text)
    run.bold = bold
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    return run


def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    add_text(title, "{{BusinessName}}", size=22, color=RGBColor(0x0B, 0x25, 0x45), bold=True)

    subtitle = doc.add_paragraph()
    add_text(subtitle, "Invoice template example", size=11, color=RGBColor(0x55, 0x55, 0x55))

    meta = doc.add_table(rows=4, cols=4)
    widths = [1860, 2820, 1860, 2820]
    set_table_geometry(meta, widths)
    meta_data = [
        ("Invoice number", "{{InvoiceNumber}}", "Issue date", "{{IssueDate}}"),
        ("Billing period", "{{StartDate}} to {{EndDate}}", "Payment terms", "{{PaymentTerms}}"),
        ("Client", "{{ClientName}}", "Client company", "{{ClientCompany}}"),
        ("Client address", "{{ClientAddress}}", "Currency", "{{Currency}}"),
    ]
    for row, values in zip(meta.rows, meta_data):
        for idx, value in enumerate(values):
            cell = row.cells[idx]
            cell.text = value
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    if idx in (0, 2):
                        run.bold = True
                        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
            if idx in (0, 2):
                set_cell_shading(cell, "F2F4F7")

    doc.add_heading("Line items", level=1)
    items = doc.add_table(rows=2, cols=6)
    item_widths = [1240, 1450, 3020, 900, 1120, 1510]
    set_table_geometry(items, item_widths)
    headers = ["Date", "Project", "Description", "Hours", "Rate", "Amount"]
    for idx, text in enumerate(headers):
        cell = items.rows[0].cells[idx]
        cell.text = text
        set_cell_shading(cell, "F2F4F7")
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
    placeholders = ["{{ItemDate}}", "{{ProjectName}}", "{{Description}}", "{{Hours}}", "{{Rate}}", "{{LineTotal}}"]
    for idx, text in enumerate(placeholders):
        cell = items.rows[1].cells[idx]
        cell.text = text
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                run.font.size = Pt(9)

    total = doc.add_paragraph()
    total.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_text(total, "Invoice total: ", bold=True, size=13)
    add_text(total, "{{Currency}} {{InvoiceTotal}}", bold=True, size=13)

    doc.add_heading("Template placeholders", level=1)
    p = doc.add_paragraph()
    p.add_run("Use ordinary Word editing to change branding, wording, fonts, and table layout. ")
    p.add_run("Keep placeholder names wrapped in double braces so Time Tracker can replace them later.").bold = True

    placeholder_table = doc.add_table(rows=1, cols=2)
    set_table_geometry(placeholder_table, [3000, 6360])
    placeholder_table.rows[0].cells[0].text = "Placeholder"
    placeholder_table.rows[0].cells[1].text = "Meaning"
    for cell in placeholder_table.rows[0].cells:
        set_cell_shading(cell, "F2F4F7")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)

    rows = [
        ("{{InvoiceNumber}}", "The immutable invoice number, for example WLSA/JN/56."),
        ("{{IssueDate}}", "The date the invoice was created."),
        ("{{ClientName}}", "The client record name."),
        ("{{ClientCompany}}", "The client company field."),
        ("{{ClientAddress}}", "The client address field."),
        ("{{LineItems}}", "A future replacement point for all invoice jobs."),
        ("{{InvoiceTotal}}", "The invoice total calculated from selected jobs."),
    ]
    for key, meaning in rows:
        row = placeholder_table.add_row()
        row.cells[0].text = key
        row.cells[1].text = meaning
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.text = "Generated template example for Time Tracker"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
